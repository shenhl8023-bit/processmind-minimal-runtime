"""
文件内容提取服务 —— 支持从 PDF / DOCX / XLSX / JSON 中提取文本
用于第二步规则提炼时，将上传文档的实际内容传给大模型分析。
"""
import json
import os
from typing import Optional


# 每个文件提取的最大字符数（避免 LLM token 超限）
MAX_CHARS_PER_FILE = 8000


def extract_text_from_pdf(filepath: str) -> str:
    """使用 PyMuPDF 提取 PDF 文本"""
    import fitz  # PyMuPDF

    text_parts = []
    try:
        doc = fitz.open(filepath)
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
    except Exception as e:
        return f"[PDF 解析失败: {e}]"

    return "\n".join(text_parts)


def extract_text_from_docx(filepath: str) -> str:
    """使用 python-docx 提取 DOCX 文本（段落 + 表格）"""
    from docx import Document as DocxDocument

    text_parts = []
    try:
        doc = DocxDocument(filepath)
        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # 提取表格
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
    except Exception as e:
        return f"[DOCX 解析失败: {e}]"

    return "\n".join(text_parts)


def extract_text_from_xlsx(filepath: str) -> str:
    """使用 openpyxl 提取 XLSX 文本"""
    from openpyxl import load_workbook

    text_parts = []
    try:
        wb = load_workbook(filepath, read_only=True, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text_parts.append(f"[工作表: {sheet_name}]")
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) for cell in row if cell is not None)
                if row_text:
                    text_parts.append(row_text)
        wb.close()
    except Exception as e:
        return f"[XLSX 解析失败: {e}]"

    return "\n".join(text_parts)


def _read_text_with_fallbacks(filepath: str, encodings: tuple[str, ...]) -> str:
    last_error: Exception | None = None
    for encoding in encodings:
        try:
            with open(filepath, "r", encoding=encoding) as fp:
                return fp.read()
        except Exception as e:
            last_error = e
    raise last_error or UnicodeDecodeError("utf-8", b"", 0, 1, "无法识别编码")


def _flatten_json_to_lines(data: object, path: list[str] | None = None) -> list[str]:
    path = path or []
    lines: list[str] = []

    if isinstance(data, dict):
        for key, value in data.items():
            key_text = str(key).strip()
            next_path = [*path, key_text] if key_text else list(path)
            if isinstance(value, (dict, list)):
                if key_text:
                    lines.append(" / ".join(next_path))
                lines.extend(_flatten_json_to_lines(value, next_path))
            else:
                if key_text:
                    lines.append(key_text)
                value_text = str(value).strip()
                if value_text:
                    lines.append(value_text)
        return lines

    if isinstance(data, list):
        for index, item in enumerate(data, start=1):
            next_path = [*path, f"item_{index}"]
            if isinstance(item, (dict, list)):
                lines.append(" / ".join(next_path))
                lines.extend(_flatten_json_to_lines(item, next_path))
            else:
                item_text = str(item).strip()
                if item_text:
                    lines.append(item_text)
        return lines

    scalar_text = str(data).strip()
    return [scalar_text] if scalar_text else []


def extract_text_from_json(filepath: str) -> str:
    """提取 JSON 文本，优先格式化结构化内容，失败时回退原文。"""
    try:
        raw_text = _read_text_with_fallbacks(filepath, ("utf-8", "utf-8-sig", "gb18030", "gbk"))
    except Exception as e:
        return f"[JSON 解析失败: {e}]"

    try:
        parsed = json.loads(raw_text)
        flattened_lines = _flatten_json_to_lines(parsed)
        if flattened_lines:
            return "\n".join(flattened_lines)
        return json.dumps(parsed, ensure_ascii=False, indent=2)
    except Exception:
        return raw_text


def extract_text(filepath: str, file_type: Optional[str] = None, max_chars: Optional[int] = MAX_CHARS_PER_FILE) -> str:
    """
    根据文件类型提取文本内容。
    
    Args:
        filepath: 文件的绝对路径
        file_type: 文件扩展名 (pdf/docx/xlsx/json)，若为空则从路径推断
    
    Returns:
        提取的文本内容（截断至 MAX_CHARS_PER_FILE）
    """
    if not os.path.exists(filepath):
        return "[文件不存在]"

    if not file_type:
        file_type = filepath.rsplit(".", 1)[-1].lower() if "." in filepath else ""

    file_type = file_type.lower()

    if file_type == "pdf":
        raw_text = extract_text_from_pdf(filepath)
    elif file_type in ("docx", "doc"):
        raw_text = extract_text_from_docx(filepath)
    elif file_type in ("xlsx", "xls"):
        raw_text = extract_text_from_xlsx(filepath)
    elif file_type == "json":
        raw_text = extract_text_from_json(filepath)
    else:
        return f"[不支持的文件类型: {file_type}]"

    # 截断过长的文本；当 max_chars=None 时返回全文，供本地抽取逻辑使用。
    if max_chars is not None and len(raw_text) > max_chars:
        raw_text = raw_text[:max_chars] + f"\n\n... (文本已截断，共 {len(raw_text)} 字符，仅显示前 {max_chars} 字符)"

    return raw_text
